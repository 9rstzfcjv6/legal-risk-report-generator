from io import BytesIO
from docx import Document

def save_markdown_report(report, file_path):
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(report)

        from io import BytesIO
from docx import Document


def clean_markdown_formatting(text):
    return (
        text
        .replace("**", "")
        .replace("__", "")
        .replace("`", "")
    )

def is_markdown_table_line(line):
    return line.startswith("|") and line.endswith("|")


def parse_markdown_table_line(line):
    cells = line.strip("|").split("|")
    return [cell.strip() for cell in cells]

def convert_markdown_to_docx_bytes(markdown_text):
    doc = Document()
    lines = markdown_text.split("\n")
    index = 0

    while index < len(lines):
        clean_line = clean_markdown_formatting(lines[index].strip())

        if not clean_line:
            doc.add_paragraph()
            index += 1
            continue

        if clean_line.startswith("# "):
            doc.add_heading(clean_line.replace("# ", ""), level=1)
            index += 1

        elif clean_line.startswith("## "):
            doc.add_heading(clean_line.replace("## ", ""), level=2)
            index += 1

        elif clean_line.startswith("### "):
            doc.add_heading(clean_line.replace("### ", ""), level=3)
            index += 1

        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line.replace("- ", ""), style="List Bullet")
            index += 1

        elif is_markdown_table_line(clean_line):
            table_lines = []

            while index < len(lines):
                table_line = clean_markdown_formatting(lines[index].strip())

                if is_markdown_table_line(table_line):
                    table_lines.append(table_line)
                    index += 1
                else:
                    break

            if len(table_lines) >= 2:
                header_cells = parse_markdown_table_line(table_lines[0])
                data_lines = table_lines[2:]

                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = "Table Grid"

                header_row = table.rows[0].cells

                for cell_index, header in enumerate(header_cells):
                    header_row[cell_index].text = header

                for data_line in data_lines:
                    data_cells = parse_markdown_table_line(data_line)
                    row_cells = table.add_row().cells

                    for cell_index, cell_value in enumerate(data_cells):
                        if cell_index < len(row_cells):
                            row_cells[cell_index].text = cell_value

                doc.add_paragraph()

            else:
                for table_line in table_lines:
                    doc.add_paragraph(table_line)

        else:
            doc.add_paragraph(clean_line)
            index += 1

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream